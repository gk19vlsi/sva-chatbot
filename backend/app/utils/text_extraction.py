"""
Text Extraction Utilities

Provides utilities for extracting text from various document formats:
- PDF files using PyMuPDF
- DOCX files using python-docx
- Markdown files using markdown library
- Plain text files

Includes optimizations for large file processing:
- Chunked processing for large files
- Streaming support
- Memory-efficient extraction
"""
import fitz  # PyMuPDF
from docx import Document
import markdown
from typing import Tuple, Generator, Optional
import logging
import os

logger = logging.getLogger(__name__)

# Configuration for chunked processing
CHUNK_SIZE = 1024 * 1024  # 1MB chunks for text processing
LARGE_FILE_THRESHOLD = 10 * 1024 * 1024  # 10MB threshold for chunked processing


def extract_text_from_pdf(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Tuple of (extracted_text, success)
    """
    try:
        doc = fitz.open(file_path)
        text_parts = []
        
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                text_parts.append(text)
        
        doc.close()
        extracted_text = "\n\n".join(text_parts)
        
        if not extracted_text.strip():
            logger.warning(f"No text extracted from PDF: {file_path}")
            return "", False
            
        return extracted_text, True
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {str(e)}")
        return "", False


def extract_text_from_docx(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Tuple of (extracted_text, success)
    """
    try:
        doc = Document(file_path)
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip():
                    text_parts.append(row_text)
        
        extracted_text = "\n\n".join(text_parts)
        
        if not extracted_text.strip():
            logger.warning(f"No text extracted from DOCX: {file_path}")
            return "", False
            
        return extracted_text, True
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {str(e)}")
        return "", False


def extract_text_from_markdown(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from Markdown file.
    
    Args:
        file_path: Path to the Markdown file
        
    Returns:
        Tuple of (extracted_text, success)
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # Return raw markdown (preserves structure better for requirement extraction)
        if not md_content.strip():
            logger.warning(f"No text extracted from Markdown: {file_path}")
            return "", False
            
        return md_content, True
        
    except Exception as e:
        logger.error(f"Error extracting text from Markdown {file_path}: {str(e)}")
        return "", False


def extract_text_from_txt(file_path: str) -> Tuple[str, bool]:
    """
    Extract text from plain text file.
    
    Args:
        file_path: Path to the text file
        
    Returns:
        Tuple of (extracted_text, success)
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        
        if not text.strip():
            logger.warning(f"No text extracted from TXT: {file_path}")
            return "", False
            
        return text, True
        
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {str(e)}")
        return "", False


def extract_text(file_path: str, file_type: str) -> Tuple[str, bool]:
    """
    Extract text from a file based on its type.
    
    Args:
        file_path: Path to the file
        file_type: File type/extension (pdf, docx, md, txt)
        
    Returns:
        Tuple of (extracted_text, success)
    """
    file_type = file_type.lower().strip('.')
    
    extractors = {
        'pdf': extract_text_from_pdf,
        'docx': extract_text_from_docx,
        'md': extract_text_from_markdown,
        'markdown': extract_text_from_markdown,
        'txt': extract_text_from_txt,
    }
    
    extractor = extractors.get(file_type)
    if not extractor:
        logger.error(f"Unsupported file type: {file_type}")
        return "", False
    
    return extractor(file_path)


def is_large_file(file_path: str) -> bool:
    """
    Check if a file is large and should use chunked processing.
    
    Args:
        file_path: Path to the file
        
    Returns:
        True if file is large, False otherwise
    """
    try:
        file_size = os.path.getsize(file_path)
        return file_size > LARGE_FILE_THRESHOLD
    except Exception as e:
        logger.error(f"Error checking file size: {e}")
        return False


def extract_text_chunked(file_path: str, file_type: str, 
                        chunk_size: int = CHUNK_SIZE) -> Generator[str, None, None]:
    """
    Extract text from a file in chunks for memory-efficient processing.
    
    Args:
        file_path: Path to the file
        file_type: File type/extension
        chunk_size: Size of each chunk in bytes
        
    Yields:
        Text chunks
        
    Validates: Requirements 1.1, 1.2, 1.3, 1.4, 2.1 - Process large files in chunks
    """
    file_type = file_type.lower().strip('.')
    
    if file_type in ['txt', 'md', 'markdown']:
        # For text-based files, read in chunks
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                while True:
                    chunk = f.read(chunk_size)
                    if not chunk:
                        break
                    yield chunk
        except Exception as e:
            logger.error(f"Error reading file in chunks: {e}")
            return
    
    elif file_type == 'pdf':
        # For PDFs, process page by page
        try:
            doc = fitz.open(file_path)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    yield text
            doc.close()
        except Exception as e:
            logger.error(f"Error extracting PDF in chunks: {e}")
            return
    
    elif file_type == 'docx':
        # For DOCX, process paragraph by paragraph
        try:
            doc = Document(file_path)
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    yield paragraph.text + "\n"
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        yield row_text + "\n"
        except Exception as e:
            logger.error(f"Error extracting DOCX in chunks: {e}")
            return
    
    else:
        logger.error(f"Unsupported file type for chunked extraction: {file_type}")
        return


async def extract_text_streaming(file_path: str, file_type: str) -> Tuple[str, bool]:
    """
    Extract text from a file using streaming for large files.
    
    This function automatically detects large files and uses chunked processing
    to avoid memory issues.
    
    Args:
        file_path: Path to the file
        file_type: File type/extension
        
    Returns:
        Tuple of (extracted_text, success)
        
    Validates: Requirements 1.1, 1.2, 1.3, 1.4 - Use streaming for large files
    """
    try:
        # Check if file is large
        if is_large_file(file_path):
            logger.info(f"Using chunked processing for large file: {file_path}")
            
            # Process in chunks and combine
            text_parts = []
            for chunk in extract_text_chunked(file_path, file_type):
                text_parts.append(chunk)
            
            extracted_text = "".join(text_parts)
            
            if not extracted_text.strip():
                logger.warning(f"No text extracted from large file: {file_path}")
                return "", False
            
            return extracted_text, True
        else:
            # Use regular extraction for small files
            return extract_text(file_path, file_type)
    
    except Exception as e:
        logger.error(f"Error in streaming text extraction: {e}")
        return "", False


def get_file_info(file_path: str) -> dict:
    """
    Get information about a file for processing decisions.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Dictionary with file information
    """
    try:
        file_size = os.path.getsize(file_path)
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_name)[1].lower().strip('.')
        
        return {
            "name": file_name,
            "size": file_size,
            "size_mb": file_size / (1024 * 1024),
            "extension": file_ext,
            "is_large": file_size > LARGE_FILE_THRESHOLD,
            "should_chunk": file_size > LARGE_FILE_THRESHOLD
        }
    except Exception as e:
        logger.error(f"Error getting file info: {e}")
        return {
            "name": "unknown",
            "size": 0,
            "size_mb": 0,
            "extension": "unknown",
            "is_large": False,
            "should_chunk": False
        }
